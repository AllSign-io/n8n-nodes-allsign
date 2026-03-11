#!/usr/bin/env python3
"""
Generate a professional NDA DOCX template with {{ variables }} for AllSign.
These variables are replaced at document creation time via the AllSign API.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = etree.SubElement(tcPr, qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')

def create_nda_template():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ══════════════════════════════════════════════════════
    # HEADER
    # ══════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(4)
    run = title.add_run('NON-DISCLOSURE AGREEMENT')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(6)
    run = subtitle.add_run('ACUERDO DE CONFIDENCIALIDAD')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    # Decorative line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.space_after = Pt(20)
    run = line.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x7b, 0xff)

    # ══════════════════════════════════════════════════════
    # PARTIES INFO TABLE
    # ══════════════════════════════════════════════════════
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 1: Date
    cell = info_table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run('Date / Fecha:')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    cell = info_table.cell(0, 1)
    p = cell.paragraphs[0]
    run = p.add_run('{{ effective_date }}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x7b, 0xff)
    run.bold = True

    # Row 2: Disclosing Party
    cell = info_table.cell(1, 0)
    p = cell.paragraphs[0]
    run = p.add_run('Disclosing Party / Parte Reveladora:')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    cell = info_table.cell(1, 1)
    p = cell.paragraphs[0]
    run = p.add_run('{{ company_name }}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x7b, 0xff)
    run.bold = True

    # Row 3: Receiving Party
    cell = info_table.cell(2, 0)
    p = cell.paragraphs[0]
    run = p.add_run('Receiving Party / Parte Receptora:')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    cell = info_table.cell(2, 1)
    p = cell.paragraphs[0]
    run = p.add_run('{{ client_name }}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x7b, 0xff)
    run.bold = True

    doc.add_paragraph()  # spacer

    # ══════════════════════════════════════════════════════
    # SECTIONS
    # ══════════════════════════════════════════════════════

    def add_section_header(text, number):
        p = doc.add_paragraph()
        p.space_before = Pt(16)
        p.space_after = Pt(6)
        run = p.add_run(f'{number}. {text}')
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    def add_body_text(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = Pt(16)
        p.space_after = Pt(8)
        return p

    # ── Section 1: Purpose ──
    add_section_header('PURPOSE / PROPÓSITO', 1)
    add_body_text(
        'This Non-Disclosure Agreement ("Agreement") is entered into as of '
        '{{ effective_date }}, by and between {{ company_name }} '
        '("Disclosing Party") and {{ client_name }} ("Receiving Party"), '
        'collectively referred to as the "Parties."'
    )
    add_body_text(
        'The purpose of this Agreement is to protect the confidential and '
        'proprietary information that may be disclosed between the Parties '
        'in connection with {{ project_description }}.'
    )

    # ── Section 2: Definition ──
    add_section_header('DEFINITION OF CONFIDENTIAL INFORMATION', 2)
    add_body_text(
        '"Confidential Information" means any data or information, oral or written, '
        'disclosed by either Party that is designated as confidential or that reasonably '
        'should be understood to be confidential given the nature of the information and '
        'the circumstances of disclosure. This includes, but is not limited to:'
    )

    items = [
        'Business plans, strategies, and financial information',
        'Technical data, trade secrets, and know-how',
        'Product plans, designs, and specifications',
        'Customer lists, marketing strategies, and sales data',
        'Software, source code, algorithms, and databases',
        'Any other information marked as "Confidential" or "Proprietary"',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.line_spacing = Pt(15)

    # ── Section 3: Obligations ──
    add_section_header('OBLIGATIONS OF RECEIVING PARTY', 3)
    add_body_text(
        'The Receiving Party agrees to:'
    )
    obligations = [
        'Hold and maintain the Confidential Information in strict confidence',
        'Not disclose the Confidential Information to any third parties without prior written consent',
        'Use the Confidential Information solely for the purpose described in Section 1',
        'Protect the Confidential Information using the same degree of care used to protect its own confidential information, but in no event less than reasonable care',
        'Promptly notify the Disclosing Party of any unauthorized disclosure or use',
    ]
    for item in obligations:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.line_spacing = Pt(15)

    # ── Section 4: Duration ──
    add_section_header('TERM AND DURATION / VIGENCIA', 4)
    add_body_text(
        'This Agreement shall remain in effect for a period of '
        '{{ confidentiality_period }} from the date of execution. '
        'The obligations of confidentiality shall survive the termination '
        'of this Agreement for an additional period of two (2) years.'
    )

    # ── Section 5: Exclusions ──
    add_section_header('EXCLUSIONS', 5)
    add_body_text(
        'Confidential Information does not include information that:'
    )
    exclusions = [
        'Is or becomes publicly available through no fault of the Receiving Party',
        'Was already known to the Receiving Party prior to disclosure',
        'Is independently developed by the Receiving Party without use of the Confidential Information',
        'Is rightfully received from a third party without restriction on disclosure',
        'Is required to be disclosed by law, regulation, or court order',
    ]
    for item in exclusions:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.line_spacing = Pt(15)

    # ── Section 6: Return of Information ──
    add_section_header('RETURN OF INFORMATION', 6)
    add_body_text(
        'Upon termination of this Agreement or upon request by the Disclosing Party, '
        'the Receiving Party shall promptly return or destroy all copies of the '
        'Confidential Information in its possession and certify in writing that it has done so.'
    )

    # ── Section 7: Governing Law ──
    add_section_header('GOVERNING LAW / LEY APLICABLE', 7)
    add_body_text(
        'This Agreement shall be governed by and construed in accordance with the laws of '
        '{{ governing_law }}. Any disputes arising under this Agreement shall be resolved '
        'in the competent courts of said jurisdiction.'
    )

    # ── Section 8: Miscellaneous ──
    add_section_header('MISCELLANEOUS', 8)
    add_body_text(
        'This Agreement constitutes the entire agreement between the Parties concerning '
        'the subject matter hereof and supersedes all prior agreements, understandings, and '
        'negotiations. No modification of this Agreement shall be effective unless in writing '
        'and signed by both Parties.'
    )

    # ══════════════════════════════════════════════════════
    # SIGNATURE BLOCK
    # ══════════════════════════════════════════════════════
    doc.add_paragraph()  # spacer

    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line2.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x7b, 0xff)

    sig_header = doc.add_paragraph()
    sig_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_header.space_after = Pt(20)
    run = sig_header.add_run('SIGNATURES / FIRMAS')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Signature table
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    for i, header in enumerate(['Disclosing Party / Parte Reveladora', 'Receiving Party / Parte Receptora']):
        cell = sig_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Company / Client names
    cell = sig_table.cell(1, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('{{ company_name }}')
    run.font.color.rgb = RGBColor(0x00, 0x7b, 0xff)
    run.bold = True

    cell = sig_table.cell(1, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('{{ client_name }}')
    run.font.color.rgb = RGBColor(0x00, 0x7b, 0xff)
    run.bold = True

    # Signature lines
    for col in range(2):
        cell = sig_table.cell(2, col)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(40)
        run = p.add_run('_' * 35)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    # "Firma del ..." labels
    cell = sig_table.cell(3, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Firma del Representante')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    cell = sig_table.cell(3, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Firma del Cliente')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    # ══════════════════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════════════════
    doc.add_paragraph()
    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_text.add_run('Document generated and signed electronically via AllSign')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
    run.italic = True

    date_footer = doc.add_paragraph()
    date_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_footer.add_run('allsign.io')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x7b, 0xff)

    # ── Save ──
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, 'NDA_Template_AllSign.docx')
    doc.save(output_path)
    print(f'✅ NDA Template created: {output_path}')
    print()
    print('Template variables used:')
    print('  {{ client_name }}           → Name of the receiving party')
    print('  {{ company_name }}          → Name of the disclosing party')
    print('  {{ effective_date }}        → Date the agreement takes effect')
    print('  {{ project_description }}   → Brief description of the project/scope')
    print('  {{ confidentiality_period }}→ Duration (e.g. "2 years")')
    print('  {{ governing_law }}         → Jurisdiction (e.g. "Mexico City, Mexico")')

if __name__ == '__main__':
    create_nda_template()
